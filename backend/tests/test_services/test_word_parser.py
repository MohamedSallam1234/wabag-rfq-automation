"""Tests for .docx → Markdown extraction (headings, paragraphs, tables)."""

from collections.abc import Callable
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType

from app.services.ingestion.word_parser import extract_docx_markdown


def _docx_path(tmp_path: Path, build: Callable[[DocxDocumentType], None]) -> str:
    document = DocxDocument()
    build(document)
    path = tmp_path / "doc.docx"
    document.save(str(path))
    return str(path)


def test_heading_paragraph_and_table(tmp_path: Path) -> None:
    def build(doc: DocxDocumentType) -> None:
        doc.add_heading("Scope of Work", level=1)
        doc.add_paragraph("Supply centrifugal pumps.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Tag"
        table.cell(0, 1).text = "Description"
        table.cell(1, 0).text = "P-101"
        table.cell(1, 1).text = "Centrifugal pump"

    markdown = extract_docx_markdown(_docx_path(tmp_path, build))

    assert "# Scope of Work" in markdown
    assert "Supply centrifugal pumps." in markdown
    assert "| Tag | Description |" in markdown
    assert "| --- | --- |" in markdown
    assert "P-101" in markdown


def test_subheading_uses_deeper_level(tmp_path: Path) -> None:
    markdown = extract_docx_markdown(
        _docx_path(tmp_path, lambda doc: doc.add_heading("Details", level=3))
    )
    assert "### Details" in markdown


def test_empty_docx_returns_empty(tmp_path: Path) -> None:
    assert extract_docx_markdown(_docx_path(tmp_path, lambda doc: None)) == ""


def test_table_cell_pipe_is_escaped(tmp_path: Path) -> None:
    def build(doc: DocxDocumentType) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "A|B"

    markdown = extract_docx_markdown(_docx_path(tmp_path, build))
    assert "A\\|B" in markdown

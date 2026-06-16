"""Tests for the document endpoints (single synchronous upload, open — no auth)."""

import io
import uuid
from collections.abc import Iterator
from datetime import UTC, datetime
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pypdf import PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle
from storage3.utils import StorageException

from app.api.deps import get_db, get_storage
from app.core.config import get_settings
from app.main import app
from app.models.document import DocTypeSource, Document, RetentionPolicy
from app.models.project import Project

_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


# --------------------------------------------------------------------------- #
# File-fixture builders (valid, parseable office files)
# --------------------------------------------------------------------------- #
def _pdf_bytes(pages: int = 1) -> bytes:
    """A textless (blank) PDF — stands in for a scanned / image-only document."""
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=72, height=72)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _text_table_pdf_bytes(pages: int = 1) -> bytes:
    """A born-digital PDF with a paragraph and a bordered table on each page."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story: list[object] = []
    for i in range(pages):
        story.append(Paragraph("Equipment specification for centrifugal pumps.", styles["Normal"]))
        rows = [["Tag", "Description", "Qty"], ["P-101", "Centrifugal pump", "2"]]
        table = Table(rows)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
        story.append(table)
        if i < pages - 1:
            story.append(PageBreak())
    doc.build(story)
    return buffer.getvalue()


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Tag", "Description"])
    sheet.append(["P-101", "Centrifugal pump"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Scope of Work", level=1)
    document.add_paragraph("Supply and install centrifugal pumps.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# Fakes
# --------------------------------------------------------------------------- #
async def _fake_refresh(instance: object, *_: object, **__: object) -> None:
    now = datetime.now(UTC)
    for attr in ("created_at", "updated_at"):
        if getattr(instance, attr, None) is None:
            setattr(instance, attr, now)


def _make_session() -> MagicMock:
    session = MagicMock()
    session.flush = AsyncMock()
    session.commit = AsyncMock()
    session.delete = AsyncMock()
    session.refresh = AsyncMock(side_effect=_fake_refresh)
    session.get = AsyncMock()
    session.execute = AsyncMock()
    return session


def _make_storage() -> tuple[MagicMock, MagicMock]:
    storage = MagicMock()
    proxy = MagicMock()
    proxy.upload = AsyncMock(return_value={"path": "p"})
    proxy.create_signed_url = AsyncMock(return_value={"signedURL": "https://download"})
    proxy.remove = AsyncMock(return_value=[])
    storage.from_ = MagicMock(return_value=proxy)
    return storage, proxy


def _make_project() -> Project:
    now = datetime.now(UTC)
    return Project(id=uuid.uuid4(), name="P", created_at=now, updated_at=now)


def _make_document(project_id: uuid.UUID, **overrides: object) -> Document:
    now = datetime.now(UTC)
    defaults: dict[str, object] = {
        "id": uuid.uuid4(),
        "project_id": project_id,
        "original_filename": "01_Spec_Rev01.pdf",
        "storage_bucket": "rfq-documents",
        "storage_path": "proj/doc.pdf",
        "content_type": "application/pdf",
        "size_bytes": 1000,
        "doc_type_source": DocTypeSource.AUTO,
        "retention": RetentionPolicy.PERSISTENT,
        "created_at": now,
        "updated_at": now,
    }
    defaults.update(overrides)
    return Document(**defaults)


@pytest.fixture(autouse=True)
def _clear_overrides() -> Iterator[None]:
    yield
    app.dependency_overrides.clear()


# --------------------------------------------------------------------------- #
# Upload (the core change)
# --------------------------------------------------------------------------- #
def test_upload_pdf_returns_201_classified_with_page_count() -> None:
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    pdf = _text_table_pdf_bytes(pages=3)
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("01_Employer_Spec_Rev02.pdf", pdf, "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    # Classification still runs on the original .pdf filename.
    assert body["doc_type"] == "Employer Technical Specifications"
    assert body["doc_type_source"] == "auto"
    assert body["revision_number"] == 2
    assert body["retention"] == "persistent"
    assert body["page_count"] == 3
    assert body["content_type"] == "text/markdown"
    proxy.upload.assert_awaited_once()
    session.add.assert_called_once()

    # The stored object is a .md artifact holding the extracted text + Markdown table.
    stored_path, stored_bytes, stored_opts = proxy.upload.call_args.args
    assert stored_path.endswith(".md")
    assert stored_opts == {"content-type": "text/markdown"}
    assert not stored_bytes.startswith(b"%PDF-")
    text = stored_bytes.decode("utf-8")
    assert "P-101" in text  # table cell preserved
    assert "---" in text  # GFM table separator


def test_upload_xlsx_returns_201_with_sheet_names() -> None:
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("03_Equipment_List.xlsx", _xlsx_bytes(), _XLSX_MIME)},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["sheet_names"] == ["Sheet"]
    assert body["doc_type"] == "Equipment List"
    assert body["retention"] == "persistent"
    assert body["content_type"] == "text/markdown"
    proxy.upload.assert_awaited_once()

    # The stored object is a .md artifact under the per-document folder, holding the sheet table.
    stored_path, stored_bytes, stored_opts = proxy.upload.call_args.args
    assert stored_path.startswith(f"projects/{project.id}/documents/")
    assert stored_path.endswith("03_Equipment_List.xlsx.md")
    assert stored_opts == {"content-type": "text/markdown"}
    text = stored_bytes.decode("utf-8")
    assert "## Sheet" in text
    assert "P-101" in text


def test_upload_docx_returns_201_md() -> None:
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("02_Scope_of_Work.docx", _docx_bytes(), "application/octet-stream")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["content_type"] == "text/markdown"
    assert body["page_count"] is None
    assert body["sheet_names"] is None
    proxy.upload.assert_awaited_once()

    stored_path, stored_bytes, stored_opts = proxy.upload.call_args.args
    assert stored_path.endswith("02_Scope_of_Work.docx.md")
    assert stored_opts == {"content-type": "text/markdown"}
    text = stored_bytes.decode("utf-8")
    assert "# Scope of Work" in text
    assert "centrifugal pumps" in text


@pytest.mark.parametrize(
    "filename",
    ["malware.exe", "notes.txt", "archive.zip", "diagram.png", "report.csv", "no_extension"],
)
def test_upload_rejects_unsupported_extension_415(filename: str) -> None:
    """Any extension outside the allow-list is rejected with 415 before anything is stored."""
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": (filename, b"some bytes", "application/octet-stream")},
    )
    assert response.status_code == 415
    proxy.upload.assert_not_awaited()  # rejected before storage
    session.add.assert_not_called()  # nothing persisted


def test_upload_content_extension_mismatch_422() -> None:
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    # XLSX bytes presented as a .pdf — the magic-byte gate rejects it.
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("report.pdf", _xlsx_bytes(), "application/pdf")},
    )
    assert response.status_code == 422
    proxy.upload.assert_not_awaited()
    session.add.assert_not_called()


def test_upload_unparseable_ooxml_422() -> None:
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    # ZIP magic clears the coarse gate, but the deep parse fails.
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("data.xlsx", b"PK\x03\x04 not really a zip", _XLSX_MIME)},
    )
    assert response.status_code == 422
    proxy.upload.assert_not_awaited()


def test_upload_huge_file_returns_413(monkeypatch: pytest.MonkeyPatch) -> None:
    """A body larger than the cap is rejected with 413; the guard trips mid-stream."""
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()
    # 1 MB cap, 2 MB payload → the running total exceeds the cap on the second chunk.
    monkeypatch.setattr(get_settings(), "MAX_UPLOAD_SIZE_MB", 1)
    payload = b"x" * (2 * 1024 * 1024)

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("big.pdf", payload, "application/pdf")},
    )
    assert response.status_code == 413
    proxy.upload.assert_not_awaited()  # nothing stored
    session.add.assert_not_called()


def test_upload_storage_failure_returns_502_and_persists_nothing() -> None:
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()
    proxy.upload = AsyncMock(side_effect=StorageException("boom"))

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    # A text+table PDF passes validation and reaches storage (a blank PDF would 422 first).
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("01_Spec_Rev01.pdf", _text_table_pdf_bytes(pages=1), "application/pdf")},
    )
    assert response.status_code == 502
    session.add.assert_not_called()  # nothing persisted when storage fails


def test_upload_scanned_pdf_returns_422_and_stores_nothing() -> None:
    """A textless (scanned / image-only) PDF is rejected 422 before anything is stored."""
    project = _make_project()
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.post(
        f"/api/v1/projects/{project.id}/documents",
        files={"file": ("scanned.pdf", _pdf_bytes(pages=1), "application/pdf")},
    )
    assert response.status_code == 422
    proxy.upload.assert_not_awaited()
    session.add.assert_not_called()


def test_upload_to_missing_project_returns_404() -> None:
    session = _make_session()
    session.get = AsyncMock(return_value=None)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.post(
        f"/api/v1/projects/{uuid.uuid4()}/documents",
        files={"file": ("01_Spec_Rev01.pdf", _pdf_bytes(pages=1), "application/pdf")},
    )
    assert response.status_code == 404
    proxy.upload.assert_not_awaited()


# --------------------------------------------------------------------------- #
# List / detail / patch / delete (de-authed)
# --------------------------------------------------------------------------- #
def test_list_documents() -> None:
    project = _make_project()
    document = _make_document(project.id)
    session = _make_session()
    session.get = AsyncMock(return_value=project)
    result = MagicMock()
    result.scalars.return_value.all.return_value = [document]
    session.execute = AsyncMock(return_value=result)

    app.dependency_overrides[get_db] = lambda: session

    client = TestClient(app)
    response = client.get(f"/api/v1/projects/{project.id}/documents")
    assert response.status_code == 200
    assert len(response.json()) == 1


def test_get_document_includes_download_url() -> None:
    project = _make_project()
    document = _make_document(project.id)
    session = _make_session()
    session.get = AsyncMock(return_value=document)
    storage, _ = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.get(f"/api/v1/documents/{document.id}")
    assert response.status_code == 200
    assert response.json()["download_url"] == "https://download"


def test_get_document_missing_returns_404() -> None:
    session = _make_session()
    session.get = AsyncMock(return_value=None)
    storage, _ = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.get(f"/api/v1/documents/{uuid.uuid4()}")
    assert response.status_code == 404


def test_get_document_signing_failure_returns_502() -> None:
    project = _make_project()
    document = _make_document(project.id)
    session = _make_session()
    session.get = AsyncMock(return_value=document)
    storage, proxy = _make_storage()
    proxy.create_signed_url = AsyncMock(side_effect=StorageException("boom"))

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.get(f"/api/v1/documents/{document.id}")
    assert response.status_code == 502


def test_get_document_empty_signed_url_returns_502() -> None:
    project = _make_project()
    document = _make_document(project.id)
    session = _make_session()
    session.get = AsyncMock(return_value=document)
    storage, proxy = _make_storage()
    proxy.create_signed_url = AsyncMock(return_value={"signedURL": ""})

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.get(f"/api/v1/documents/{document.id}")
    assert response.status_code == 502


def test_patch_document_classification_sets_manual_and_recomputes_retention() -> None:
    project = _make_project()
    document = _make_document(project.id, retention=RetentionPolicy.TRANSIENT)
    session = _make_session()
    session.get = AsyncMock(return_value=document)

    app.dependency_overrides[get_db] = lambda: session

    client = TestClient(app)
    response = client.patch(f"/api/v1/documents/{document.id}", json={"doc_type": "Equipment List"})
    assert response.status_code == 200
    body = response.json()
    assert body["doc_type"] == "Equipment List"
    assert body["doc_type_source"] == "manual"
    assert body["retention"] == "persistent"  # Equipment List is a project-common doc


def test_delete_document_removes_object() -> None:
    project = _make_project()
    document = _make_document(project.id)
    session = _make_session()
    session.get = AsyncMock(return_value=document)
    storage, proxy = _make_storage()

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.delete(f"/api/v1/documents/{document.id}")
    assert response.status_code == 204
    session.delete.assert_awaited_once()
    proxy.remove.assert_awaited_once()


def test_delete_document_swallows_storage_error() -> None:
    project = _make_project()
    document = _make_document(project.id)
    session = _make_session()
    session.get = AsyncMock(return_value=document)
    storage, proxy = _make_storage()
    proxy.remove = AsyncMock(side_effect=StorageException("gone"))

    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_storage] = lambda: storage

    client = TestClient(app)
    response = client.delete(f"/api/v1/documents/{document.id}")
    assert response.status_code == 204  # storage orphan is logged, not fatal
    session.delete.assert_awaited_once()
